# app/main.py
import os
import sys
import tempfile
import subprocess
from pathlib import Path
from typing import List, Optional
import psycopg2
from psycopg2.extras import RealDictCursor, execute_values
from fastapi import FastAPI, Request, Form, UploadFile, File
from fastapi.responses import HTMLResponse, RedirectResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from docxtpl import DocxTemplate
from docx import Document

# Add parent directory to path so we can import modules from the workspace
sys.path.append(str(Path(__file__).resolve().parent.parent))
import app.importer as importer
import diagnostics

app = FastAPI(title="DClinPsy Viva Matching Console")

# Mount static and templates
BASE_DIR = Path(__file__).resolve().parent
app.mount("/static", StaticFiles(directory=str(BASE_DIR / "static")), name="static")
templates = Jinja2Templates(directory=str(BASE_DIR / "templates"))

DB_DSN = "dbname=vivas user=vivas_admin password=treefrog host=localhost port=5432"

# Temporary cache for uploaded survey files to be committed
class UploadCache:
    file_internal: Optional[Path] = None
    file_external: Optional[Path] = None
    file_thesis: Optional[Path] = None

upload_cache = UploadCache()

def get_db_conn():
    return psycopg2.connect(DB_DSN)

@app.on_event("startup")
def startup_event():
    # 1. Initialize tables and system settings
    conn = get_db_conn()
    with conn, conn.cursor() as cur:
        # Create system_settings table if it doesn't exist
        cur.execute("""
            CREATE TABLE IF NOT EXISTS system_settings (
                key TEXT PRIMARY KEY,
                value TEXT
            );
        """)
        # Initialize default paths
        defaults = {
            "path_trainee": "templates/Trainee_Initial_Email.docx",
            "path_internal": "templates/Internal_Examiner_Initial_Email.docx",
            "path_triad": "templates/External_Examiner_Initial_Email.docx"
        }
        for k, v in defaults.items():
            cur.execute("""
                INSERT INTO system_settings (key, value) VALUES (%s, %s)
                ON CONFLICT (key) DO NOTHING;
            """, (k, v))
            
        # Ensure student_supervisors table exists
        cur.execute("""
            CREATE TABLE IF NOT EXISTS student_supervisors (
                student_id   INT PRIMARY KEY REFERENCES students(student_id),
                supervisors  TEXT
            );
        """)

        # Ensure students table has pre_matched column
        cur.execute("""
            ALTER TABLE students ADD COLUMN IF NOT EXISTS pre_matched BOOLEAN NOT NULL DEFAULT FALSE;
        """)
        cur.execute("ALTER TABLE students ADD COLUMN IF NOT EXISTS email TEXT;")
        cur.execute("ALTER TABLE students ADD COLUMN IF NOT EXISTS project_title TEXT;")
        cur.execute("ALTER TABLE students ADD COLUMN IF NOT EXISTS sample_other_desc TEXT;")
        cur.execute("ALTER TABLE students ADD COLUMN IF NOT EXISTS methodology_other_desc TEXT;")
        cur.execute("ALTER TABLE students ADD COLUMN IF NOT EXISTS field_other_desc TEXT;")
        cur.execute("ALTER TABLE students ADD COLUMN IF NOT EXISTS additional_characteristics TEXT;")
        cur.execute("ALTER TABLE examiners ADD COLUMN IF NOT EXISTS email TEXT;")
        cur.execute("ALTER TABLE examiners ADD COLUMN IF NOT EXISTS has_dclinpsy BOOLEAN NOT NULL DEFAULT TRUE;")
        
        # Ensure pending_mappings table exists
        cur.execute("""
            CREATE TABLE IF NOT EXISTS pending_mappings (
                mapping_id          SERIAL PRIMARY KEY,
                student_id          INT NOT NULL REFERENCES students(student_id) ON DELETE CASCADE,
                relationship_type   TEXT NOT NULL CHECK (relationship_type IN ('supervisor', 'partner')),
                raw_name            TEXT NOT NULL,
                suggested_id        INT,
                confidence          TEXT NOT NULL CHECK (confidence IN ('exact', 'fuzzy', 'none')),
                status              TEXT NOT NULL DEFAULT 'pending' CHECK (status IN ('pending', 'confirmed', 'rejected'))
            );
        """)

        # Ensure project_groups table exists
        cur.execute("""
            CREATE TABLE IF NOT EXISTS project_groups (
                group_id   INT NOT NULL,
                student_id INT UNIQUE REFERENCES students(student_id) ON DELETE CASCADE,
                PRIMARY KEY (group_id, student_id)
            );
        """)
        
        # Drop group size trigger/function if it exists (no longer enforced)
        cur.execute("DROP TRIGGER IF EXISTS trg_project_group_size ON project_groups")
        cur.execute("DROP FUNCTION IF EXISTS check_group_size()")
        
        # Ensure examiner_limits table allows limit_n >= 0
        cur.execute("ALTER TABLE examiner_limits DROP CONSTRAINT IF EXISTS examiner_limits_limit_n_check;")
        cur.execute("ALTER TABLE examiner_limits ADD CONSTRAINT examiner_limits_limit_n_check CHECK (limit_n >= 0);")
        
        # 2. Fill missing examiner limits
        cur.execute("""
            INSERT INTO examiner_limits (examiner_id, limit_n)
            SELECT examiner_id, 3
            FROM examiners
            ON CONFLICT (examiner_id) DO NOTHING;
        """)
    conn.close()

# Redirect root to Tab 1
@app.get("/")
def read_root():
    return RedirectResponse(url="/import")

# ===========================================================================
# TAB 1: DATA IMPORT
# ===========================================================================
@app.get("/import", response_class=HTMLResponse)
def get_import(request: Request):
    return templates.TemplateResponse(request, "import_tab.html", {"active_tab": "import"})

@app.post("/import/validate-survey")
async def validate_survey(
    request: Request,
    file_internal: UploadFile = File(...),
    file_external: UploadFile = File(...),
    file_thesis: UploadFile = File(...)
):
    # Save files to temp directory
    temp_dir = Path(tempfile.gettempdir())
    
    # Write files
    path_int = temp_dir / "temp_internal.xlsx"
    path_ext = temp_dir / "temp_external.xlsx"
    path_proj = temp_dir / "temp_thesis.xlsx"
    
    with path_int.open("wb") as f:
        f.write(await file_internal.read())
    with path_ext.open("wb") as f:
        f.write(await file_external.read())
    with path_proj.open("wb") as f:
        f.write(await file_thesis.read())
        
    # Validate
    ok, msg, stats = importer.validate_survey_files(path_int, path_ext, path_proj)
    
    if ok:
        # Cache paths for commit
        upload_cache.file_internal = path_int
        upload_cache.file_external = path_ext
        upload_cache.file_thesis = path_proj
        
        html = f"""
        <div class="alert alert-success">
            <strong>Verification Successful!</strong><br>
            Found {stats['student_count']} students, {stats['internal_count']} internal examiners, 
            {stats['external_count']} external examiners, and {stats['categories_count']} competence categories.
        </div>
        <div class="form-group" style="margin-bottom: 12px; max-width: 320px;">
            <label style="font-weight: 500; font-size: 0.85rem; display: block; margin-bottom: 6px;">Import Mode</label>
            <select name="import_mode" id="import-mode-select" style="background-color: var(--bg-secondary); border: 1px solid var(--border); color: var(--text-primary); border-radius: var(--radius-sm); padding: 8px; width: 100%; font-size: 0.9rem;">
                <option value="add">Add / Merge Records (Iterative Build)</option>
                <option value="overwrite">Overwrite Database (Clear Slate)</option>
            </select>
        </div>
        <button hx-post="/import/commit-survey" hx-include="#import-mode-select" hx-target="#survey-validation-result" hx-indicator="#commit-indicator" class="btn btn-primary">
            Commit to Database
        </button>
        <span id="commit-indicator" class="htmx-indicator" style="margin-left: 10px;">
            <span class="spinner"></span> Committing to Postgres...
        </span>
        """
    else:
        # Delete temp files immediately on failure
        for p in [path_int, path_ext, path_proj]:
            if p.exists():
                p.unlink()
        html = f"""
        <div class="alert alert-danger">
            <strong>Validation Failed:</strong><br>
            <pre style="white-space: pre-wrap; font-family: monospace; font-size: 0.85rem; margin-top: 8px;">{msg}</pre>
        </div>
        """
    return HTMLResponse(content=html)

@app.post("/import/commit-survey")
def commit_survey(import_mode: str = Form("add")):
    if not (upload_cache.file_internal and upload_cache.file_external and upload_cache.file_thesis):
        return HTMLResponse(content="""
            <div class="alert alert-danger">
                <strong>Error:</strong> Uploaded file cache expired. Please re-validate files first.
            </div>
        """)
        
    overwrite = (import_mode == "overwrite")
    ok, msg = importer.commit_survey_to_db(
        DB_DSN,
        upload_cache.file_internal,
        upload_cache.file_external,
        upload_cache.file_thesis,
        overwrite=overwrite
    )
    
    # Cleanup files
    for p in [upload_cache.file_internal, upload_cache.file_external, upload_cache.file_thesis]:
        if p.exists():
            p.unlink()
    upload_cache.file_internal = None
    upload_cache.file_external = None
    upload_cache.file_thesis = None
    
    # Replenish missing examiner limits
    conn = get_db_conn()
    with conn, conn.cursor() as cur:
        cur.execute("""
            INSERT INTO examiner_limits (examiner_id, limit_n)
            SELECT examiner_id, 3
            FROM examiners
            ON CONFLICT (examiner_id) DO NOTHING;
        """)
    conn.close()
    
    if ok:
        html = f"""
        <div class="alert alert-success">
            <strong>Success:</strong> Database tables successfully populated!
        </div>
        """
        return HTMLResponse(content=html, headers={"HX-Trigger": "refreshVerificationCount, surveyCommitted, refreshVerificationCard"})
    else:
        html = f"""
        <div class="alert alert-danger">
            <strong>Database Commit Failed:</strong><br>
            <pre style="white-space: pre-wrap; font-family: monospace; font-size: 0.85rem; margin-top: 8px;">{msg}</pre>
        </div>
        """
        return HTMLResponse(content=html)


@app.post("/import/update-competences")
async def update_competences(
    file_internal: UploadFile = File(None),
    file_external: UploadFile = File(None)
):
    """Re-import competence spreadsheets to add/update examiner competences.
    
    This only touches examiners and examiner_competences tables.
    Students, thesis categories, pending mappings, bans, groups, and 
    assignments are left completely untouched.
    """
    if (not file_internal or not file_internal.filename) and (not file_external or not file_external.filename):
        return HTMLResponse(content='<div class="alert alert-danger">Please upload at least one competence spreadsheet (Internal or External).</div>')

    import tempfile
    upload_dir = Path("uploads")
    upload_dir.mkdir(exist_ok=True)
    
    path_int = None
    path_ext = None
    
    if file_internal and file_internal.filename:
        path_int = upload_dir / f"reupload_internal_{file_internal.filename}"
        with path_int.open("wb") as f:
            f.write(await file_internal.read())
            
    if file_external and file_external.filename:
        path_ext = upload_dir / f"reupload_external_{file_external.filename}"
        with path_ext.open("wb") as f:
            f.write(await file_external.read())
    
    try:
        df_int = None
        df_ext = None
        if path_int:
            df_int = importer.tidy_columns(importer.read_sheet(path_int, key_col_name="examiner"), 'examiner')
        if path_ext:
            df_ext = importer.tidy_columns(importer.read_sheet(path_ext, key_col_name="examiner"), 'examiner')
        
        competences = []
        if df_int is not None and df_ext is not None:
            competences_int = sorted(set(df_int.columns) - {'examiner'})
            competences_ext = sorted(set(df_ext.columns) - {'examiner'})
            if set(competences_int) != set(competences_ext):
                mismatch = []
                if set(competences_int) - set(competences_ext):
                    mismatch.append(f"Only in Internal: {sorted(set(competences_int) - set(competences_ext))}")
                if set(competences_ext) - set(competences_int):
                    mismatch.append(f"Only in External: {sorted(set(competences_ext) - set(competences_int))}")
                html = f"""
                <div class="alert alert-danger">
                    <strong>Column Mismatch:</strong><br>
                    <pre style="white-space: pre-wrap; font-size: 0.85rem; margin-top: 8px;">{"<br>".join(mismatch)}</pre>
                </div>
                """
                return HTMLResponse(content=html)
            competences = competences_int
        elif df_int is not None:
            competences = sorted(set(df_int.columns) - {'examiner'})
        elif df_ext is not None:
            competences = sorted(set(df_ext.columns) - {'examiner'})
        
        # Validate
        if df_int is not None:
            ok_int, msg_int = importer.validate_examiner_df(df_int, competences)
            if not ok_int:
                html = f"""
                <div class="alert alert-danger">
                    <strong>Internal Sheet Error:</strong><br>
                    <pre style="white-space: pre-wrap; font-size: 0.85rem; margin-top: 8px;">{msg_int}</pre>
                </div>
                """
                return HTMLResponse(content=html)
        
        if df_ext is not None:
            ok_ext, msg_ext = importer.validate_examiner_df(df_ext, competences)
            if not ok_ext:
                html = f"""
                <div class="alert alert-danger">
                    <strong>External Sheet Error:</strong><br>
                    <pre style="white-space: pre-wrap; font-size: 0.85rem; margin-top: 8px;">{msg_ext}</pre>
                </div>
                """
                return HTMLResponse(content=html)
        
        # Commit
        conn = get_db_conn()
        conn.autocommit = False
        try:
            with conn, conn.cursor() as cur:
                # Upsert categories
                execute_values(
                    cur,
                    "INSERT INTO categories (name) VALUES %s ON CONFLICT (name) DO NOTHING",
                    [(c,) for c in competences]
                )
                cur.execute("SELECT category_id, name FROM categories")
                cat_id = {n: i for i, n in cur.fetchall()}
                
                # Upsert examiners and competences
                map_level = {'can': 'can', 'could': 'could', "can't": 'cannot'}
                
                new_examiners = []
                updated_examiners = []
                
                datasets = []
                if df_int is not None:
                    datasets.append((df_int, 'internal'))
                if df_ext is not None:
                    datasets.append((df_ext, 'external'))
                
                for df, e_type in datasets:
                    for row in df.itertuples():
                        # Check if this examiner already exists
                        cur.execute("SELECT examiner_id FROM examiners WHERE full_name = %s", (row.examiner,))
                        existing = cur.fetchone()
                        
                        cur.execute("""
                            INSERT INTO examiners (full_name, examiner_type)
                            VALUES (%s, %s)
                            ON CONFLICT (full_name) DO UPDATE SET examiner_type = EXCLUDED.examiner_type
                            RETURNING examiner_id
                        """, (row.examiner, e_type))
                        eid = cur.fetchone()[0]
                        
                        if existing:
                            updated_examiners.append(row.examiner)
                        else:
                            new_examiners.append(row.examiner)
                        
                        # Upsert competences
                        tuples = []
                        for comp in competences:
                            level = map_level[str(getattr(row, comp)).strip().lower()]
                            tuples.append((eid, cat_id[comp], level))
                        
                        execute_values(
                            cur,
                            "INSERT INTO examiner_competences (examiner_id, category_id, competence) VALUES %s "
                            "ON CONFLICT (examiner_id, category_id) DO UPDATE SET competence = EXCLUDED.competence",
                            tuples
                        )
                
                # Fill missing examiner limits for new examiners
                cur.execute("""
                    INSERT INTO examiner_limits (examiner_id, limit_n)
                    SELECT examiner_id, 3
                    FROM examiners
                    ON CONFLICT (examiner_id) DO NOTHING;
                """)
                
                # Update any pending_mappings that had no match but now might resolve
                cur.execute("SELECT examiner_id, full_name FROM examiners")
                all_examiners = {r[1].strip(): r[0] for r in cur.fetchall()}
                
                cur.execute("""
                    SELECT mapping_id, raw_name 
                    FROM pending_mappings 
                    WHERE relationship_type = 'supervisor' 
                    AND status = 'pending' 
                    AND (suggested_id IS NULL OR confidence = 'none')
                """)
                unresolved = cur.fetchall()
                auto_resolved = 0
                for mapping_id, raw_name in unresolved:
                    match = importer.match_name(raw_name, all_examiners)
                    if match:
                        suggested_eid, suggested_name, confidence = match
                        cur.execute("""
                            UPDATE pending_mappings 
                            SET suggested_id = %s, confidence = %s 
                            WHERE mapping_id = %s
                        """, (suggested_eid, confidence, mapping_id))
                        auto_resolved += 1
                
            conn.commit()
            
            summary_parts = []
            if new_examiners:
                summary_parts.append(f"<strong>{len(new_examiners)} new examiner(s)</strong> added: {', '.join(new_examiners[:5])}{'...' if len(new_examiners) > 5 else ''}")
            if updated_examiners:
                summary_parts.append(f"<strong>{len(updated_examiners)} existing examiner(s)</strong> updated with refreshed competences")
            if auto_resolved:
                summary_parts.append(f"<strong>{auto_resolved} pending supervisor mapping(s)</strong> auto-resolved to newly available examiners")
            
            html = f"""
            <div class="alert alert-success">
                <strong>Competences Updated Successfully!</strong><br>
                {'<br>'.join(summary_parts)}
            </div>
            """
            return HTMLResponse(content=html, headers={"HX-Trigger": "refreshVerificationCount, refreshVerificationCard"})
        except Exception as e:
            conn.rollback()
            import traceback
            traceback.print_exc()
            html = f"""
            <div class="alert alert-danger">
                <strong>Database Error:</strong><br>
                <pre style="white-space: pre-wrap; font-size: 0.85rem; margin-top: 8px;">{str(e)}</pre>
            </div>
            """
            return HTMLResponse(content=html)
        finally:
            conn.close()
    finally:
        for p in [path_int, path_ext]:
            if p and p.exists():
                p.unlink()

@app.post("/import/reset-db")
def reset_db():
    conn = get_db_conn()
    try:
        with conn, conn.cursor() as cur:
            cur.execute("DELETE FROM pending_mappings;")
            cur.execute("DELETE FROM allocation_audit;")
            cur.execute("DELETE FROM examiner_assignments;")
            cur.execute("DELETE FROM examiner_student_bans;")
            cur.execute("DELETE FROM student_supervisors;")
            cur.execute("DELETE FROM project_groups;")
            cur.execute("DELETE FROM thesis_categories;")
            cur.execute("DELETE FROM examiner_competences;")
            cur.execute("DELETE FROM examiner_limits;")
            cur.execute("DELETE FROM students;")
            cur.execute("DELETE FROM examiners;")
            cur.execute("DELETE FROM categories;")
        conn.commit()
        return HTMLResponse(content="", headers={"HX-Refresh": "true"})
    except Exception as e:
        conn.rollback()
        html = f'<div class="alert alert-danger" style="margin-top: 10px;"><strong>Reset Failed:</strong> {str(e)}</div>'
        return HTMLResponse(content=html)
    finally:
        conn.close()


@app.get("/import/verification-dashboard")
def get_verification_dashboard(request: Request, active_tab: str = None, errors: List[str] = None):
    conn = get_db_conn()
    try:
        with conn, conn.cursor() as cur:
            # Fetch all pending mappings
            cur.execute("""
                SELECT 
                    pm.mapping_id, 
                    s.full_name AS student_name, 
                    pm.relationship_type,
                    pm.raw_name, 
                    pm.suggested_id, 
                    pm.confidence
                FROM pending_mappings pm
                JOIN students s USING (student_id)
                WHERE pm.status = 'pending'
                ORDER BY pm.relationship_type, s.full_name, pm.raw_name
            """)
            exact_sups = []
            exact_parts = []
            fuzzy_sups = []
            fuzzy_parts = []
            none_sups = []
            none_parts = []
            for r in cur.fetchall():
                item = {
                    "mapping_id": r[0],
                    "student_name": r[1],
                    "type": r[2], # 'supervisor' or 'partner'
                    "raw_name": r[3],
                    "suggested_id": r[4],
                    "confidence": r[5]
                }
                if r[2] == 'supervisor':
                    if r[5] == 'exact':
                        exact_sups.append(item)
                    elif r[5] == 'fuzzy':
                        fuzzy_sups.append(item)
                    else:
                        none_sups.append(item)
                elif r[2] == 'partner':
                    if r[5] == 'exact':
                        exact_parts.append(item)
                    elif r[5] == 'fuzzy':
                        fuzzy_parts.append(item)
                    else:
                        none_parts.append(item)

            # Fetch examiners for selection
            cur.execute("SELECT examiner_id, full_name, examiner_type FROM examiners ORDER BY full_name")
            examiners = [{"id": r[0], "name": r[1], "type": r[2]} for r in cur.fetchall()]

            # Fetch students for selection
            cur.execute("SELECT student_id, full_name FROM students ORDER BY full_name")
            all_students = [{"id": r[0], "name": r[1]} for r in cur.fetchall()]

            total_pending_supervisors = len(exact_sups) + len(fuzzy_sups) + len(none_sups)
            total_pending_partners = len(exact_parts) + len(fuzzy_parts) + len(none_parts)
            total_pending = total_pending_supervisors + total_pending_partners

            # Determine default active tab if not specified
            if active_tab is None:
                if total_pending_supervisors > 0:
                    active_tab = "supervisor"
                elif total_pending_partners > 0:
                    active_tab = "partner"
                else:
                    active_tab = "supervisor"

        return templates.TemplateResponse(request, "verification_dashboard.html", {
            "exact_sups": exact_sups,
            "exact_parts": exact_parts,
            "fuzzy_sups": fuzzy_sups,
            "fuzzy_parts": fuzzy_parts,
            "none_sups": none_sups,
            "none_parts": none_parts,
            "total_pending_supervisors": total_pending_supervisors,
            "total_pending_partners": total_pending_partners,
            "examiners": examiners,
            "students": all_students,
            "total_pending": total_pending,
            "active_tab": active_tab,
            "errors": errors
        })
    except Exception as e:
        import traceback
        traceback.print_exc()
        return HTMLResponse(content=f"Error loading dashboard: {str(e)}", status_code=500)
    finally:
        conn.close()


@app.get("/import/verification-count")
def get_verification_count():
    conn = get_db_conn()
    try:
        with conn, conn.cursor() as cur:
            cur.execute("SELECT COUNT(*) FROM pending_mappings WHERE status = 'pending'")
            count = cur.fetchone()[0]
        # Return count as HTML badge
        if count > 0:
            return HTMLResponse(content=f'<span class="badge badge-warning" style="margin-left: 5px;">{count}</span>')
        return HTMLResponse(content="")
    except Exception as e:
        return HTMLResponse(content="")
    finally:
        conn.close()


@app.post("/import/verification/update-suggestion")
def update_suggestion(
    mapping_id: int = Form(...),
    selected_id: str = Form(None)
):
    conn = get_db_conn()
    try:
        with conn, conn.cursor() as cur:
            sel_id = int(selected_id) if selected_id and selected_id.strip() else None
            cur.execute("""
                UPDATE pending_mappings
                SET suggested_id = %s
                WHERE mapping_id = %s
            """, (sel_id, mapping_id))
        conn.commit()
        return HTMLResponse(content="")
    except Exception as e:
        return HTMLResponse(content=f"Error saving selection: {str(e)}", status_code=500)
    finally:
        conn.close()


@app.post("/import/verification/confirm")
def confirm_verification(
    mapping_id: int = Form(...),
    selected_id: str = Form(None)
):
    conn = get_db_conn()
    try:
        with conn, conn.cursor() as cur:
            cur.execute("SELECT student_id, relationship_type, raw_name FROM pending_mappings WHERE mapping_id = %s", (mapping_id,))
            row = cur.fetchone()
            if not row:
                return HTMLResponse(content="Mapping not found", status_code=404)
            
            sid, r_type, raw_name = row
            sel_id = int(selected_id) if selected_id and selected_id.strip() else None

            if r_type == 'supervisor':
                if sel_id:
                    cur.execute("""
                        INSERT INTO examiner_student_bans (examiner_id, student_id, reason)
                        VALUES (%s, %s, %s)
                        ON CONFLICT (examiner_id, student_id) DO NOTHING
                    """, (sel_id, sid, f"Project Supervisor ({raw_name})"))
            
            elif r_type == 'partner':
                if sel_id:
                    # Fetch group members
                    cur.execute("""
                        SELECT student_id FROM project_groups 
                        WHERE group_id = (SELECT group_id FROM project_groups WHERE student_id = %s)
                    """, (sid,))
                    g1_members = {r[0] for r in cur.fetchall()} or {sid}

                    cur.execute("""
                        SELECT student_id FROM project_groups 
                        WHERE group_id = (SELECT group_id FROM project_groups WHERE student_id = %s)
                    """, (sel_id,))
                    g2_members = {r[0] for r in cur.fetchall()} or {sel_id}

                    cur.execute("SELECT group_id FROM project_groups WHERE student_id = %s", (sid,))
                    g1 = cur.fetchone()
                    cur.execute("SELECT group_id FROM project_groups WHERE student_id = %s", (sel_id,))
                    g2 = cur.fetchone()
                    
                    if g1 and g2:
                        g1_id = g1[0]
                        g2_id = g2[0]
                        if g1_id != g2_id:
                            cur.execute("UPDATE project_groups SET group_id = %s WHERE group_id = %s", (g1_id, g2_id))
                    elif g1:
                        cur.execute("INSERT INTO project_groups (group_id, student_id) VALUES (%s, %s)", (g1[0], sel_id))
                    elif g2:
                        cur.execute("INSERT INTO project_groups (group_id, student_id) VALUES (%s, %s)", (g2[0], sid))
                    else:
                        cur.execute("SELECT COALESCE(MAX(group_id), 0) + 1 FROM project_groups")
                        new_gid = cur.fetchone()[0]
                        cur.execute("INSERT INTO project_groups (group_id, student_id) VALUES (%s, %s)", (new_gid, sid))
                        cur.execute("INSERT INTO project_groups (group_id, student_id) VALUES (%s, %s)", (new_gid, sel_id))
            
            cur.execute("""
                UPDATE pending_mappings
                SET status = 'confirmed', suggested_id = %s
                WHERE mapping_id = %s
            """, (sel_id, mapping_id))

        conn.commit()
        return HTMLResponse(content="", headers={"HX-Trigger": "refreshVerificationCount"})
    except Exception as e:
        return HTMLResponse(content=f"Error confirming mapping: {str(e)}", status_code=500)
    finally:
        conn.close()


@app.post("/import/verification/reject")
def reject_verification(mapping_id: int = Form(...)):
    conn = get_db_conn()
    try:
        with conn, conn.cursor() as cur:
            cur.execute("""
                UPDATE pending_mappings
                SET status = 'rejected'
                WHERE mapping_id = %s
            """, (mapping_id,))
        conn.commit()
        return HTMLResponse(content="", headers={"HX-Trigger": "refreshVerificationCount"})
    except Exception as e:
        return HTMLResponse(content=f"Error rejecting mapping: {str(e)}", status_code=500)
    finally:
        conn.close()


def confirm_all_by_confidence(request: Request, conf: str, relationship_type: str):
    conn = get_db_conn()
    errors = []
    try:
        with conn, conn.cursor() as cur:
            if conf == 'none':
                cur.execute("""
                    SELECT mapping_id, student_id, relationship_type, raw_name, suggested_id 
                    FROM pending_mappings 
                    WHERE status = 'pending' AND confidence = 'none' AND relationship_type = %s
                """, (relationship_type,))
            else:
                cur.execute("""
                    SELECT mapping_id, student_id, relationship_type, raw_name, suggested_id 
                    FROM pending_mappings 
                    WHERE status = 'pending' AND confidence = %s AND relationship_type = %s
                """, (conf, relationship_type))
            rows = cur.fetchall()
            for r in rows:
                mapping_id, sid, r_type, raw_name, sel_id = r
                
                cur.execute(f"SAVEPOINT sp_{mapping_id}")
                try:
                    if r_type == 'supervisor':
                        if sel_id:
                            cur.execute("""
                                INSERT INTO examiner_student_bans (examiner_id, student_id, reason)
                                VALUES (%s, %s, %s)
                                ON CONFLICT (examiner_id, student_id) DO NOTHING
                            """, (sel_id, sid, f"Project Supervisor ({raw_name})"))
                    
                    elif r_type == 'partner':
                        if sel_id:
                            cur.execute("""
                                SELECT student_id FROM project_groups 
                                WHERE group_id = (SELECT group_id FROM project_groups WHERE student_id = %s)
                            """, (sid,))
                            g1_members = {row[0] for row in cur.fetchall()} or {sid}

                            cur.execute("""
                                SELECT student_id FROM project_groups 
                                WHERE group_id = (SELECT group_id FROM project_groups WHERE student_id = %s)
                            """, (sel_id,))
                            g2_members = {row[0] for row in cur.fetchall()} or {sel_id}

                            cur.execute("SELECT group_id FROM project_groups WHERE student_id = %s", (sid,))
                            g1 = cur.fetchone()
                            cur.execute("SELECT group_id FROM project_groups WHERE student_id = %s", (sel_id,))
                            g2 = cur.fetchone()
                            
                            if g1 and g2:
                                g1_id = g1[0]
                                g2_id = g2[0]
                                if g1_id != g2_id:
                                    cur.execute("UPDATE project_groups SET group_id = %s WHERE group_id = %s", (g1_id, g2_id))
                            elif g1:
                                cur.execute("INSERT INTO project_groups (group_id, student_id) VALUES (%s, %s)", (g1[0], sel_id))
                            elif g2:
                                cur.execute("INSERT INTO project_groups (group_id, student_id) VALUES (%s, %s)", (g2[0], sid))
                            else:
                                cur.execute("SELECT COALESCE(MAX(group_id), 0) + 1 FROM project_groups")
                                new_gid = cur.fetchone()[0]
                                cur.execute("INSERT INTO project_groups (group_id, student_id) VALUES (%s, %s)", (new_gid, sid))
                                cur.execute("INSERT INTO project_groups (group_id, student_id) VALUES (%s, %s)", (new_gid, sel_id))
                    
                    cur.execute("""
                        UPDATE pending_mappings 
                        SET status = 'confirmed' 
                        WHERE mapping_id = %s
                    """, (mapping_id,))
                    cur.execute(f"RELEASE SAVEPOINT sp_{mapping_id}")
                except Exception as e:
                    cur.execute(f"ROLLBACK TO SAVEPOINT sp_{mapping_id}")
                    cur.execute("SELECT full_name FROM students WHERE student_id = %s", (sid,))
                    st_name = cur.fetchone()[0]
                    errors.append(f"Student '{st_name}': {str(e)}")
                
        conn.commit()
        return get_verification_dashboard(request, active_tab=relationship_type, errors=errors)
    except Exception as e:
        import traceback
        traceback.print_exc()
        return HTMLResponse(content=f"Error: {str(e)}", status_code=500)
    finally:
        conn.close()


@app.post("/import/verification/confirm-all-exact")
def confirm_all_exact(request: Request, relationship_type: str = Form(...)):
    return confirm_all_by_confidence(request, 'exact', relationship_type)


@app.post("/import/verification/confirm-all-fuzzy")
def confirm_all_fuzzy(request: Request, relationship_type: str = Form(...)):
    return confirm_all_by_confidence(request, 'fuzzy', relationship_type)


@app.post("/import/verification/confirm-all-no-match")
def confirm_all_no_match(request: Request, relationship_type: str = Form(...)):
    return confirm_all_by_confidence(request, 'none', relationship_type)


@app.post("/import/verification/create-examiner")
def create_examiner(
    request: Request,
    mapping_id: int = Form(...),
    examiner_type: str = Form("external")
):
    conn = get_db_conn()
    try:
        with conn, conn.cursor() as cur:
            cur.execute("SELECT raw_name FROM pending_mappings WHERE mapping_id = %s", (mapping_id,))
            row = cur.fetchone()
            if not row:
                return HTMLResponse(content="Mapping not found", status_code=404)
            
            raw_name = row[0]
            # 1. Create examiner
            cur.execute("""
                INSERT INTO examiners (full_name, examiner_type)
                VALUES (%s, %s)
                ON CONFLICT (full_name) DO UPDATE SET examiner_type = EXCLUDED.examiner_type
                RETURNING examiner_id
            """, (raw_name, examiner_type))
            eid = cur.fetchone()[0]

            # Initialize limit for the new examiner
            cur.execute("""
                INSERT INTO examiner_limits (examiner_id, limit_n)
                VALUES (%s, 3)
                ON CONFLICT DO NOTHING
            """, (eid,))

            # Initialize all competences to 'could' for this new examiner
            cur.execute("SELECT category_id FROM categories")
            cat_ids = [r[0] for r in cur.fetchall()]
            for cid in cat_ids:
                cur.execute("""
                    INSERT INTO examiner_competences (examiner_id, category_id, competence)
                    VALUES (%s, %s, 'could')
                    ON CONFLICT (examiner_id, category_id) DO NOTHING
                """, (eid, cid))
            
            # 2. Update mapping to point to the new examiner and set confidence to exact
            cur.execute("""
                UPDATE pending_mappings
                SET suggested_id = %s, confidence = 'exact'
                WHERE mapping_id = %s
            """, (eid, mapping_id))
            
        conn.commit()
        response = get_verification_dashboard(request, active_tab="supervisor")
        response.headers["HX-Trigger"] = "refreshVerificationCount"
        return response
    except Exception as e:
        return HTMLResponse(content=f"Error: {str(e)}", status_code=500)
    finally:
        conn.close()

@app.post("/import/export-templates")
def export_templates():
    try:
        # Run export_contacts.py as a subprocess to keep the same environment structure
        cmd = [sys.executable, "export_contacts.py", "--dsn", DB_DSN]
        res = subprocess.run(cmd, capture_output=True, text=True)
        if res.returncode == 0:
            html = f'<div class="alert alert-success"><strong>Success:</strong> {res.stdout.strip()}</div>'
        else:
            html = f'<div class="alert alert-danger"><strong>Error:</strong> {res.stderr.strip() or res.stdout.strip()}</div>'
    except Exception as e:
        html = f'<div class="alert alert-danger"><strong>Exception:</strong> {str(e)}</div>'
    return HTMLResponse(content=html)


# ===========================================================================
# TAB 2: CONSTRAINT EDITING
# ===========================================================================
def fetch_tab2_data():
    conn = get_db_conn()
    with conn, conn.cursor(cursor_factory=RealDictCursor) as cur:
        # Categories
        cur.execute("SELECT category_id AS id, name, default_weight FROM categories ORDER BY name")
        categories = cur.fetchall()
        
        # Examiners
        cur.execute("""
            SELECT ex.examiner_id, ex.full_name, ex.examiner_type, ex.has_dclinpsy, COALESCE(el.limit_n, 3) AS limit_n
            FROM examiners ex
            LEFT JOIN examiner_limits el USING (examiner_id)
            ORDER BY ex.full_name
        """)
        examiners = cur.fetchall()
        
        # Students
        cur.execute("SELECT student_id, full_name FROM students ORDER BY full_name")
        students = cur.fetchall()
        
        # Bans list
        cur.execute("""
            SELECT b.student_id, s.full_name AS student_name, b.examiner_id, e.full_name AS examiner_name, b.reason
            FROM examiner_student_bans b
            JOIN students s USING (student_id)
            JOIN examiners e USING (examiner_id)
            ORDER BY s.full_name, e.full_name
        """)
        bans = cur.fetchall()
        
        # Project groups list
        cur.execute("""
            SELECT g.group_id, string_agg(s.full_name, ', ' ORDER BY s.full_name) AS members
            FROM project_groups g
            JOIN students s USING (student_id)
            GROUP BY g.group_id
            ORDER BY g.group_id
        """)
        groups = cur.fetchall()
        # Pre-matched list
        cur.execute("""
            SELECT s.student_id, s.full_name AS student_name,
                   MAX(CASE WHEN ea.role = 'internal' THEN ex.full_name END) AS internal_name,
                   MAX(CASE WHEN ea.role = 'external' THEN ex.full_name END) AS external_name
            FROM students s
            LEFT JOIN examiner_assignments ea USING (student_id)
            LEFT JOIN examiners ex ON ea.examiner_id = ex.examiner_id
            WHERE s.pre_matched = TRUE
            GROUP BY s.student_id, s.full_name
            ORDER BY s.full_name
        """)
        prematches = cur.fetchall()
    conn.close()
    return categories, examiners, students, bans, groups, prematches

@app.get("/constraints", response_class=HTMLResponse)
def get_constraints(request: Request):
    categories, examiners, students, bans, groups, prematches = fetch_tab2_data()
    return templates.TemplateResponse(request, "constraints_tab.html", {
        "active_tab": "constraints",
        "categories": categories,
        "examiners": examiners,
        "students": students,
        "bans": bans,
        "groups": groups,
        "prematches": prematches
    })

@app.post("/constraints/save-weights")
async def save_weights(request: Request):
    form = await request.form()
    conn = get_db_conn()
    updated = 0
    try:
        with conn, conn.cursor() as cur:
            for k, v in form.items():
                if k.startswith("weight_"):
                    cat_id = int(k.replace("weight_", ""))
                    weight_val = float(v)
                    cur.execute("UPDATE categories SET default_weight = %s WHERE category_id = %s", (weight_val, cat_id))
                    updated += 1
        conn.commit()
        return HTMLResponse(content=f'<div class="alert alert-success" style="margin-top: 10px;">Successfully updated {updated} competence weights.</div>')
    except Exception as e:
        conn.rollback()
        return HTMLResponse(content=f'<div class="alert alert-danger" style="margin-top: 10px;">Error updating weights: {str(e)}</div>')
    finally:
        conn.close()

@app.post("/constraints/save-limits")
async def save_limits(request: Request):
    form = await request.form()
    conn = get_db_conn()
    updated = 0
    try:
        with conn, conn.cursor() as cur:
            for k, v in form.items():
                if k.startswith("limit_"):
                    ex_id = int(k.replace("limit_", ""))
                    limit_val = int(v)
                    cur.execute("""
                        INSERT INTO examiner_limits (examiner_id, limit_n) VALUES (%s, %s)
                        ON CONFLICT (examiner_id) DO UPDATE SET limit_n = EXCLUDED.limit_n;
                    """, (ex_id, limit_val))
                    updated += 1
        conn.commit()
        return HTMLResponse(content=f'<div class="alert alert-success" style="margin-top: 10px;">Successfully updated {updated} examiner capacity limits.</div>')
    except Exception as e:
        conn.rollback()
        return HTMLResponse(content=f'<div class="alert alert-danger" style="margin-top: 10px;">Error updating capacity limits: {str(e)}</div>')
    finally:
        conn.close()

@app.post("/constraints/add-ban")
async def add_ban(request: Request, student_id: int = Form(...), examiner_id: int = Form(...), reason: str = Form("")):
    conn = get_db_conn()
    with conn, conn.cursor() as cur:
        cur.execute("""
            INSERT INTO examiner_student_bans (student_id, examiner_id, reason)
            VALUES (%s, %s, %s)
            ON CONFLICT (student_id, examiner_id) DO UPDATE SET reason = EXCLUDED.reason;
        """, (student_id, examiner_id, reason.strip() or None))
    conn.commit()
    conn.close()
    
    # Re-render list
    _, _, _, bans, _, _ = fetch_tab2_data()
    return templates.TemplateResponse(request, "bans_list.html", {"bans": bans})

@app.post("/constraints/remove-ban")
async def remove_ban(request: Request, student_id: int = Form(...), examiner_id: int = Form(...)):
    conn = get_db_conn()
    with conn, conn.cursor() as cur:
        cur.execute("DELETE FROM examiner_student_bans WHERE student_id = %s AND examiner_id = %s", (student_id, examiner_id))
    conn.commit()
    conn.close()
    
    _, _, _, bans, _, _ = fetch_tab2_data()
    return templates.TemplateResponse(request, "bans_list.html", {"bans": bans})

@app.post("/constraints/add-prematch")
async def add_prematch(
    request: Request,
    student_id: int = Form(...),
    internal_id: Optional[int] = Form(None),
    external_id: Optional[int] = Form(None)
):
    conn = get_db_conn()
    alert_msg = ""
    alert_ok = True
    
    if not internal_id and not external_id:
        alert_ok = False
        alert_msg = "FAILED: You must choose at least one internal or external examiner to pre-match."
    else:
        with conn, conn.cursor() as cur:
            # 1. Capacity checks for internal
            if internal_id:
                cur.execute("SELECT COALESCE(limit_n, 3) FROM examiner_limits WHERE examiner_id = %s", (internal_id,))
                limit_row = cur.fetchone()
                limit_n = limit_row[0] if limit_row else 3
                
                cur.execute("SELECT COUNT(*) FROM examiner_assignments WHERE examiner_id = %s AND student_id != %s", (internal_id, student_id))
                curr_assigned = cur.fetchone()[0]
                
                if curr_assigned + 1 > limit_n:
                    alert_ok = False
                    cur.execute("SELECT full_name FROM examiners WHERE examiner_id = %s", (internal_id,))
                    ex_name = cur.fetchone()[0]
                    alert_msg = f"FAILED: Internal examiner {ex_name} has reached their workload limit of {limit_n} vivas."
            
            # 2. Capacity checks for external
            if alert_ok and external_id:
                cur.execute("SELECT COALESCE(limit_n, 3) FROM examiner_limits WHERE examiner_id = %s", (external_id,))
                limit_row = cur.fetchone()
                limit_n = limit_row[0] if limit_row else 3
                
                cur.execute("SELECT COUNT(*) FROM examiner_assignments WHERE examiner_id = %s AND student_id != %s", (external_id, student_id))
                curr_assigned = cur.fetchone()[0]
                
                if curr_assigned + 1 > limit_n:
                    alert_ok = False
                    cur.execute("SELECT full_name FROM examiners WHERE examiner_id = %s", (external_id,))
                    ex_name = cur.fetchone()[0]
                    alert_msg = f"FAILED: External examiner {ex_name} has reached their workload limit of {limit_n} vivas."
            
            # 3. Apply changes if validation passes
            if alert_ok:
                cur.execute("UPDATE students SET pre_matched = TRUE WHERE student_id = %s", (student_id,))
                cur.execute("DELETE FROM examiner_assignments WHERE student_id = %s", (student_id,))
                if internal_id:
                    cur.execute("INSERT INTO examiner_assignments (student_id, role, examiner_id) VALUES (%s, 'internal', %s)", (student_id, internal_id))
                if external_id:
                    cur.execute("INSERT INTO examiner_assignments (student_id, role, examiner_id) VALUES (%s, 'external', %s)", (student_id, external_id))
                
                cur.execute("SELECT full_name FROM students WHERE student_id = %s", (student_id,))
                st_name = cur.fetchone()[0]
                alert_msg = f"Pre-match saved for {st_name} successfully."
        conn.commit()
    conn.close()
    
    # Re-render pre-matches list
    _, _, _, _, _, prematches = fetch_tab2_data()
    return templates.TemplateResponse(request, "prematches_list.html", {
        "prematches": prematches,
        "alert_message": alert_msg,
        "alert_success": alert_ok
    })

@app.post("/constraints/remove-prematch")
async def remove_prematch(request: Request, student_id: int = Form(...)):
    conn = get_db_conn()
    alert_msg = ""
    with conn, conn.cursor() as cur:
        cur.execute("UPDATE students SET pre_matched = FALSE WHERE student_id = %s", (student_id,))
        cur.execute("DELETE FROM examiner_assignments WHERE student_id = %s", (student_id,))
        cur.execute("SELECT full_name FROM students WHERE student_id = %s", (student_id,))
        st_name = cur.fetchone()[0]
        alert_msg = f"Removed pre-match and reset student pool status for {st_name}."
    conn.commit()
    conn.close()
    
    # Re-render pre-matches list
    _, _, _, _, _, prematches = fetch_tab2_data()
    return templates.TemplateResponse(request, "prematches_list.html", {
        "prematches": prematches,
        "alert_message": alert_msg,
        "alert_success": True
    })

@app.post("/constraints/save-group")
async def save_group(request: Request):
    form = await request.form()
    student_ids_raw = form.getlist("student_ids")
    student_ids = [int(sid) for sid in student_ids_raw]
    group_dest = form.get("group_dest")
    existing_group_id = form.get("existing_group_id")
    
    if not student_ids:
        _, _, _, _, groups, _ = fetch_tab2_data()
        return templates.TemplateResponse(request, "groups_list.html", {"groups": groups})
        
    conn = get_db_conn()
    with conn, conn.cursor() as cur:
        # Determine group_id
        if group_dest == "existing" and existing_group_id:
            group_id = int(existing_group_id)
        else:
            cur.execute("SELECT COALESCE(MAX(group_id), 0) + 1 FROM project_groups")
            group_id = cur.fetchone()[0]
            
        # Delete existing entries for these students
        cur.execute("DELETE FROM project_groups WHERE student_id = ANY(%s)", (student_ids,))
        
        # Insert new links
        for sid in student_ids:
            cur.execute("INSERT INTO project_groups (group_id, student_id) VALUES (%s, %s) ON CONFLICT DO NOTHING", (group_id, sid))
            
    conn.commit()
    conn.close()
    
    _, _, _, _, groups, _ = fetch_tab2_data()
    return templates.TemplateResponse(request, "groups_list.html", {"groups": groups})

@app.post("/constraints/remove-group")
async def remove_group(request: Request, group_id: int = Form(...)):
    conn = get_db_conn()
    with conn, conn.cursor() as cur:
        cur.execute("DELETE FROM project_groups WHERE group_id = %s", (group_id,))
    conn.commit()
    conn.close()
    
    _, _, _, _, groups, _ = fetch_tab2_data()
    return templates.TemplateResponse(request, "groups_list.html", {"groups": groups})


# ===========================================================================
# TAB 2.5: PERSON DATA MANAGEMENT
# ===========================================================================
@app.get("/person-data", response_class=HTMLResponse)
def get_person_data(request: Request):
    return templates.TemplateResponse(request, "person_data_tab.html", {"active_tab": "person-data"})


@app.get("/person-data/list", response_class=HTMLResponse)
def get_person_data_list(request: Request, type: str = "students", search: Optional[str] = None):
    conn = get_db_conn()
    try:
        with conn, conn.cursor(cursor_factory=RealDictCursor) as cur:
            if type == "students":
                if search and search.strip():
                    q = f"%{search.strip().lower()}%"
                    cur.execute("""
                        SELECT student_id, full_name, project_title, deferred, pre_matched
                        FROM students
                        WHERE LOWER(full_name) LIKE %s OR LOWER(COALESCE(project_title, '')) LIKE %s
                        ORDER BY full_name
                    """, (q, q))
                else:
                    cur.execute("""
                        SELECT student_id, full_name, project_title, deferred, pre_matched
                        FROM students
                        ORDER BY full_name
                    """)
                students = cur.fetchall()
                return templates.TemplateResponse(request, "person_data_students_list.html", {"students": students})
            else:
                if search and search.strip():
                    q = f"%{search.strip().lower()}%"
                    cur.execute("""
                        SELECT ex.examiner_id, ex.full_name, ex.examiner_type, ex.email, ex.has_dclinpsy, COALESCE(el.limit_n, 3) AS limit_n,
                               (SELECT COUNT(*) FROM examiner_assignments WHERE examiner_id = ex.examiner_id) AS active_assignments
                        FROM examiners ex
                        LEFT JOIN examiner_limits el USING (examiner_id)
                        WHERE LOWER(ex.full_name) LIKE %s OR LOWER(COALESCE(ex.email, '')) LIKE %s
                        ORDER BY ex.full_name
                    """, (q, q))
                else:
                    cur.execute("""
                        SELECT ex.examiner_id, ex.full_name, ex.examiner_type, ex.email, ex.has_dclinpsy, COALESCE(el.limit_n, 3) AS limit_n,
                               (SELECT COUNT(*) FROM examiner_assignments WHERE examiner_id = ex.examiner_id) AS active_assignments
                        FROM examiners ex
                        LEFT JOIN examiner_limits el USING (examiner_id)
                        ORDER BY ex.full_name
                    """)
                examiners = cur.fetchall()
                return templates.TemplateResponse(request, "person_data_examiners_list.html", {"examiners": examiners})
    except Exception as e:
        return HTMLResponse(content=f'<div class="alert alert-danger">Error querying list: {str(e)}</div>')
    finally:
        conn.close()


@app.get("/person-data/student/{student_id}", response_class=HTMLResponse)
def get_student_edit(request: Request, student_id: int):
    conn = get_db_conn()
    try:
        with conn, conn.cursor(cursor_factory=RealDictCursor) as cur:
            cur.execute("""
                SELECT student_id, full_name, email, project_title, deferred, pre_matched,
                       sample_other_desc, methodology_other_desc, field_other_desc, additional_characteristics, created_at
                FROM students
                WHERE student_id = %s
            """, (student_id,))
            student = cur.fetchone()
            if not student:
                return HTMLResponse(content='<div class="alert alert-danger">Student not found.</div>')
            
            # Fetch supervisors
            cur.execute("""
                SELECT supervisors
                FROM student_supervisors
                WHERE student_id = %s
            """, (student_id,))
            sup_row = cur.fetchone()
            student['supervisors'] = sup_row['supervisors'] if sup_row else ""
            
            # Fetch partners
            cur.execute("""
                SELECT s.student_id, s.full_name
                FROM students s
                JOIN project_groups pg ON s.student_id = pg.student_id
                WHERE pg.group_id = (
                    SELECT group_id FROM project_groups WHERE student_id = %s
                ) AND s.student_id != %s
            """, (student_id, student_id))
            student['partners'] = cur.fetchall()
            
            cur.execute("""
                SELECT c.category_id, c.name, c.default_weight, tc.in_scope, tc.weight AS custom_weight
                FROM categories c
                LEFT JOIN thesis_categories tc ON c.category_id = tc.category_id AND tc.student_id = %s
                ORDER BY c.name
            """, (student_id,))
            categories = cur.fetchall()
            
        return templates.TemplateResponse(request, "student_edit_form.html", {
            "student": student,
            "categories": categories,
            "active_edit_tab": "basic"
        })
    except Exception as e:
        return HTMLResponse(content=f'<div class="alert alert-danger">Error fetching student: {str(e)}</div>')
    finally:
        conn.close()


@app.post("/person-data/student/{student_id}", response_class=HTMLResponse)
async def post_student_edit(
    request: Request,
    student_id: int,
    full_name: str = Form(...),
    email: Optional[str] = Form(None),
    project_title: Optional[str] = Form(None),
    deferred: Optional[str] = Form(None),
    sample_other_desc: Optional[str] = Form(None),
    methodology_other_desc: Optional[str] = Form(None),
    field_other_desc: Optional[str] = Form(None),
    additional_characteristics: Optional[str] = Form(None),
    supervisors: Optional[str] = Form(None)
):
    conn = get_db_conn()
    alert_message = ""
    alert_success = True
    is_deferred = (deferred == "on")
    
    form_data = await request.form()
    active_edit_tab = form_data.get("active_edit_tab", "basic")
    category_ids = [int(val) for name, val in form_data.items() if name == "category_ids"]
    
    try:
        with conn, conn.cursor() as cur:
            cur.execute("SELECT student_id FROM students WHERE LOWER(full_name) = LOWER(%s) AND student_id != %s", (full_name.strip(), student_id))
            if cur.fetchone():
                alert_message = f"Error: A student named '{full_name.strip()}' already exists."
                alert_success = False
            else:
                cur.execute("""
                    UPDATE students
                    SET full_name = %s,
                        email = %s,
                        project_title = %s,
                        deferred = %s,
                        sample_other_desc = %s,
                        methodology_other_desc = %s,
                        field_other_desc = %s,
                        additional_characteristics = %s
                    WHERE student_id = %s
                """, (
                    full_name.strip(),
                    email.strip() if email and email.strip() else None,
                    project_title.strip() if project_title and project_title.strip() else None,
                    is_deferred,
                    sample_other_desc.strip() if sample_other_desc and sample_other_desc.strip() else None,
                    methodology_other_desc.strip() if methodology_other_desc and methodology_other_desc.strip() else None,
                    field_other_desc.strip() if field_other_desc and field_other_desc.strip() else None,
                    additional_characteristics.strip() if additional_characteristics and additional_characteristics.strip() else None,
                    student_id
                ))
                
                # Update student supervisors
                supervisors_val = supervisors.strip() if supervisors else ""
                if supervisors_val:
                    cur.execute("""
                        INSERT INTO student_supervisors (student_id, supervisors)
                        VALUES (%s, %s)
                        ON CONFLICT (student_id)
                        DO UPDATE SET supervisors = EXCLUDED.supervisors
                    """, (student_id, supervisors_val))
                else:
                    cur.execute("DELETE FROM student_supervisors WHERE student_id = %s", (student_id,))
                
                # Update thesis_categories
                cur.execute("DELETE FROM thesis_categories WHERE student_id = %s", (student_id,))
                
                cur.execute("SELECT category_id FROM categories")
                all_cat_ids = [r[0] for r in cur.fetchall()]
                
                for cid in all_cat_ids:
                    in_scope = (cid in category_ids)
                    weight_val = None
                    if in_scope:
                        w_str = form_data.get(f"weight_{cid}")
                        if w_str and w_str.strip():
                            try:
                                weight_val = float(w_str.strip())
                            except ValueError:
                                pass
                    cur.execute("""
                        INSERT INTO thesis_categories (student_id, category_id, in_scope, weight)
                        VALUES (%s, %s, %s, %s)
                    """, (student_id, cid, in_scope, weight_val))
                
                conn.commit()
                alert_message = "Changes saved successfully."
                alert_success = True
    except Exception as e:
        conn.rollback()
        alert_message = f"Error saving changes: {str(e)}"
        alert_success = False
    
    try:
        with conn, conn.cursor(cursor_factory=RealDictCursor) as cur:
            cur.execute("SELECT * FROM students WHERE student_id = %s", (student_id,))
            student = cur.fetchone()
            
            # Fetch supervisors
            cur.execute("""
                SELECT supervisors
                FROM student_supervisors
                WHERE student_id = %s
            """, (student_id,))
            sup_row = cur.fetchone()
            student['supervisors'] = sup_row['supervisors'] if sup_row else ""
            
            # Fetch partners
            cur.execute("""
                SELECT s.student_id, s.full_name
                FROM students s
                JOIN project_groups pg ON s.student_id = pg.student_id
                WHERE pg.group_id = (
                    SELECT group_id FROM project_groups WHERE student_id = %s
                ) AND s.student_id != %s
            """, (student_id, student_id))
            student['partners'] = cur.fetchall()
            
            cur.execute("""
                SELECT c.category_id, c.name, c.default_weight, tc.in_scope, tc.weight AS custom_weight
                FROM categories c
                LEFT JOIN thesis_categories tc ON c.category_id = tc.category_id AND tc.student_id = %s
                ORDER BY c.name
            """, (student_id,))
            categories = cur.fetchall()
        
        headers = {"HX-Trigger": "refreshList"} if alert_success else {}
        return templates.TemplateResponse(
            request, 
            "student_edit_form.html", 
            {
                "student": student, 
                "categories": categories,
                "active_edit_tab": active_edit_tab,
                "alert_message": alert_message, 
                "alert_success": alert_success
            },
            headers=headers
        )
    except Exception as e:
        return HTMLResponse(content=f'<div class="alert alert-danger">Error rendering form: {str(e)}</div>')
    finally:
        conn.close()


@app.get("/person-data/examiner/{examiner_id}", response_class=HTMLResponse)
def get_examiner_edit(request: Request, examiner_id: int):
    conn = get_db_conn()
    try:
        with conn, conn.cursor(cursor_factory=RealDictCursor) as cur:
            cur.execute("""
                SELECT ex.examiner_id, ex.full_name, ex.email, ex.examiner_type, ex.has_dclinpsy, ex.created_at, COALESCE(el.limit_n, 3) AS limit_n,
                       (SELECT COUNT(*) FROM examiner_assignments WHERE examiner_id = ex.examiner_id) AS active_assignments
                FROM examiners ex
                LEFT JOIN examiner_limits el USING (examiner_id)
                WHERE ex.examiner_id = %s
            """, (examiner_id,))
            examiner = cur.fetchone()
            if not examiner:
                return HTMLResponse(content='<div class="alert alert-danger">Examiner not found.</div>')
            
            cur.execute("""
                SELECT c.category_id, c.name, ec.competence
                FROM categories c
                LEFT JOIN examiner_competences ec ON c.category_id = ec.category_id AND ec.examiner_id = %s
                ORDER BY c.name
            """, (examiner_id,))
            categories = cur.fetchall()
            
        return templates.TemplateResponse(request, "examiner_edit_form.html", {
            "examiner": examiner,
            "categories": categories,
            "active_edit_tab": "basic"
        })
    except Exception as e:
        return HTMLResponse(content=f'<div class="alert alert-danger">Error fetching examiner: {str(e)}</div>')
    finally:
        conn.close()


@app.post("/person-data/examiner/{examiner_id}", response_class=HTMLResponse)
async def post_examiner_edit(
    request: Request,
    examiner_id: int,
    full_name: str = Form(...),
    email: Optional[str] = Form(None),
    examiner_type: str = Form(...),
    limit_n: int = Form(...)
):
    conn = get_db_conn()
    alert_message = ""
    alert_success = True
    
    form_data = await request.form()
    active_edit_tab = form_data.get("active_edit_tab", "basic")
    has_dclinpsy_bool = (form_data.get("has_dclinpsy") == "yes")
    
    try:
        with conn, conn.cursor() as cur:
            cur.execute("SELECT examiner_id FROM examiners WHERE LOWER(full_name) = LOWER(%s) AND examiner_id != %s", (full_name.strip(), examiner_id))
            if cur.fetchone():
                alert_message = f"Error: An examiner named '{full_name.strip()}' already exists."
                alert_success = False
            else:
                cur.execute("SELECT COUNT(*) FROM examiner_assignments WHERE examiner_id = %s", (examiner_id,))
                active_assignments = cur.fetchone()[0]
                
                if limit_n < active_assignments:
                    alert_message = f"Error: Cannot reduce workload limit to {limit_n} as this examiner has {active_assignments} active viva assignments."
                    alert_success = False
                else:
                    cur.execute("""
                        UPDATE examiners
                        SET full_name = %s,
                            email = %s,
                            examiner_type = %s,
                            has_dclinpsy = %s
                        WHERE examiner_id = %s
                    """, (full_name.strip(), email.strip() if email and email.strip() else None, examiner_type, has_dclinpsy_bool, examiner_id))
                    
                    cur.execute("""
                        INSERT INTO examiner_limits (examiner_id, limit_n)
                        VALUES (%s, %s)
                        ON CONFLICT (examiner_id) DO UPDATE SET limit_n = EXCLUDED.limit_n
                    """, (examiner_id, limit_n))
                    
                    # Update examiner_competences
                    cur.execute("SELECT category_id FROM categories")
                    all_cat_ids = [r[0] for r in cur.fetchall()]
                    
                    for cid in all_cat_ids:
                        comp_val = form_data.get(f"competence_{cid}", "cannot").strip().lower()
                        if comp_val not in {"can", "could", "cannot"}:
                            comp_val = "cannot"
                        cur.execute("""
                            INSERT INTO examiner_competences (examiner_id, category_id, competence)
                            VALUES (%s, %s, %s)
                            ON CONFLICT (examiner_id, category_id) DO UPDATE SET competence = EXCLUDED.competence
                        """, (examiner_id, cid, comp_val))
                    
                    conn.commit()
                    alert_message = "Changes saved successfully."
                    alert_success = True
    except Exception as e:
        conn.rollback()
        alert_message = f"Error saving changes: {str(e)}"
        alert_success = False
        
    try:
        with conn, conn.cursor(cursor_factory=RealDictCursor) as cur:
            cur.execute("""
                SELECT ex.examiner_id, ex.full_name, ex.email, ex.examiner_type, ex.has_dclinpsy, ex.created_at, COALESCE(el.limit_n, 3) AS limit_n,
                       (SELECT COUNT(*) FROM examiner_assignments WHERE examiner_id = ex.examiner_id) AS active_assignments
                FROM examiners ex
                LEFT JOIN examiner_limits el USING (examiner_id)
                WHERE ex.examiner_id = %s
            """, (examiner_id,))
            examiner = cur.fetchone()
            
            cur.execute("""
                SELECT c.category_id, c.name, ec.competence
                FROM categories c
                LEFT JOIN examiner_competences ec ON c.category_id = ec.category_id AND ec.examiner_id = %s
                ORDER BY c.name
            """, (examiner_id,))
            categories = cur.fetchall()
            
        headers = {"HX-Trigger": "refreshList"} if alert_success else {}
        return templates.TemplateResponse(
            request, 
            "examiner_edit_form.html", 
            {
                "examiner": examiner, 
                "categories": categories,
                "active_edit_tab": active_edit_tab,
                "alert_message": alert_message, 
                "alert_success": alert_success
            },
            headers=headers
        )
    except Exception as e:
        return HTMLResponse(content=f'<div class="alert alert-danger">Error rendering form: {str(e)}</div>')
    finally:
        conn.close()


@app.post("/person-data/quick-add-examiner", response_class=HTMLResponse)
def post_quick_add_examiner(
    request: Request,
    full_name: str = Form(...),
    examiner_type: str = Form(...)
):
    full_name_clean = full_name.strip()
    if not full_name_clean:
        return HTMLResponse(content='<div class="alert alert-danger" style="margin: 20px;">Name cannot be empty.</div>')
    
    conn = get_db_conn()
    try:
        with conn, conn.cursor(cursor_factory=RealDictCursor) as cur:
            # Check if examiner with that name already exists
            cur.execute("SELECT examiner_id FROM examiners WHERE LOWER(full_name) = LOWER(%s)", (full_name_clean,))
            duplicate = cur.fetchone()
            if duplicate:
                ex_id = duplicate['examiner_id']
                cur.execute("""
                    SELECT ex.examiner_id, ex.full_name, ex.email, ex.examiner_type, ex.has_dclinpsy, ex.created_at, COALESCE(el.limit_n, 3) AS limit_n,
                           (SELECT COUNT(*) FROM examiner_assignments WHERE examiner_id = ex.examiner_id) AS active_assignments
                    FROM examiners ex
                    LEFT JOIN examiner_limits el USING (examiner_id)
                    WHERE ex.examiner_id = %s
                """, (ex_id,))
                examiner = cur.fetchone()
                
                cur.execute("""
                    SELECT c.category_id, c.name, ec.competence
                    FROM categories c
                    LEFT JOIN examiner_competences ec ON c.category_id = ec.category_id AND ec.examiner_id = %s
                    ORDER BY c.name
                """, (ex_id,))
                categories = cur.fetchall()
                
                return templates.TemplateResponse(
                    request, 
                    "examiner_edit_form.html", 
                    {
                        "examiner": examiner, 
                        "categories": categories,
                        "active_edit_tab": "basic",
                        "alert_message": f"Error: An examiner named '{full_name_clean}' already exists.", 
                        "alert_success": False
                    }
                )

            # Insert new examiner
            cur.execute("""
                INSERT INTO examiners (full_name, examiner_type)
                VALUES (%s, %s)
                RETURNING examiner_id
            """, (full_name_clean, examiner_type))
            examiner_id = cur.fetchone()['examiner_id']

            # Initialize limit to 3
            cur.execute("""
                INSERT INTO examiner_limits (examiner_id, limit_n)
                VALUES (%s, 3)
                ON CONFLICT (examiner_id) DO UPDATE SET limit_n = EXCLUDED.limit_n
            """, (examiner_id,))

            # Initialize all competences to 'could'
            cur.execute("SELECT category_id FROM categories")
            cat_ids = [r['category_id'] for r in cur.fetchall()]
            for cid in cat_ids:
                cur.execute("""
                    INSERT INTO examiner_competences (examiner_id, category_id, competence)
                    VALUES (%s, %s, 'could')
                    ON CONFLICT (examiner_id, category_id) DO NOTHING
                """, (examiner_id, cid))

            # Fetch the newly created examiner
            cur.execute("""
                SELECT ex.examiner_id, ex.full_name, ex.email, ex.examiner_type, ex.has_dclinpsy, ex.created_at, COALESCE(el.limit_n, 3) AS limit_n,
                       (SELECT COUNT(*) FROM examiner_assignments WHERE examiner_id = ex.examiner_id) AS active_assignments
                FROM examiners ex
                LEFT JOIN examiner_limits el USING (examiner_id)
                WHERE ex.examiner_id = %s
            """, (examiner_id,))
            examiner = cur.fetchone()

            cur.execute("""
                SELECT c.category_id, c.name, ec.competence
                FROM categories c
                LEFT JOIN examiner_competences ec ON c.category_id = ec.category_id AND ec.examiner_id = %s
                ORDER BY c.name
            """, (examiner_id,))
            categories = cur.fetchall()

        headers = {"HX-Trigger": "refreshList"}
        return templates.TemplateResponse(
            request, 
            "examiner_edit_form.html", 
            {
                "examiner": examiner, 
                "categories": categories,
                "active_edit_tab": "basic",
                "alert_message": f"Successfully created new examiner '{full_name_clean}'.", 
                "alert_success": True
            },
            headers=headers
        )
    except Exception as e:
        return HTMLResponse(content=f'<div class="alert alert-danger" style="margin: 20px;">Error creating examiner: {str(e)}</div>')
    finally:
        conn.close()


@app.post("/person-data/delete-student", response_class=HTMLResponse)
async def delete_student(student_id: int = Form(...)):
    conn = get_db_conn()
    alert_message = ""
    alert_success = True
    try:
        with conn, conn.cursor() as cur:
            cur.execute("DELETE FROM student_supervisors WHERE student_id = %s", (student_id,))
            cur.execute("DELETE FROM pending_mappings WHERE student_id = %s", (student_id,))
            cur.execute("DELETE FROM examiner_student_bans WHERE student_id = %s", (student_id,))
            cur.execute("DELETE FROM examiner_assignments WHERE student_id = %s", (student_id,))
            cur.execute("DELETE FROM project_groups WHERE student_id = %s", (student_id,))
            cur.execute("DELETE FROM thesis_categories WHERE student_id = %s", (student_id,))
            cur.execute("DELETE FROM allocation_audit WHERE student_id = %s", (student_id,))
            cur.execute("DELETE FROM students WHERE student_id = %s", (student_id,))
        conn.commit()
        alert_message = "Student successfully deleted."
        alert_success = True
    except Exception as e:
        conn.rollback()
        alert_message = f"Error deleting student: {str(e)}"
        alert_success = False
        
    if alert_success:
        html = f"""
        <div class="card" style="margin-bottom: 0;">
            <div class="alert alert-success" style="padding: 10px 14px; font-size: 0.85rem; margin-bottom: 0;">
                {alert_message}
            </div>
            <div style="padding: 40px 20px; text-align: center; color: var(--text-muted); border: 1px dashed var(--border); border-radius: var(--radius-md); background-color: rgba(255,255,255,0.01); margin-top: 20px;">
                <div style="font-size: 2rem; margin-bottom: 12px;">👤</div>
                <div style="font-weight: 500; color: var(--text-secondary);">No Person Selected</div>
                <p style="font-size: 0.85rem; margin-top: 4px;">Click any row on the left to view and edit details.</p>
            </div>
        </div>
        """
        return HTMLResponse(content=html, headers={"HX-Trigger": "refreshList"})
    else:
        return HTMLResponse(content=f'<div class="alert alert-danger">{alert_message}</div>')


@app.post("/person-data/delete-examiner", response_class=HTMLResponse)
async def delete_examiner(examiner_id: int = Form(...)):
    conn = get_db_conn()
    alert_message = ""
    alert_success = True
    try:
        with conn, conn.cursor() as cur:
            cur.execute("DELETE FROM examiner_limits WHERE examiner_id = %s", (examiner_id,))
            cur.execute("DELETE FROM examiner_competences WHERE examiner_id = %s", (examiner_id,))
            cur.execute("DELETE FROM examiner_student_bans WHERE examiner_id = %s", (examiner_id,))
            cur.execute("DELETE FROM examiner_assignments WHERE examiner_id = %s", (examiner_id,))
            cur.execute("DELETE FROM allocation_audit WHERE internal_examiner_id = %s OR external_examiner_id = %s", (examiner_id, examiner_id))
            cur.execute("UPDATE pending_mappings SET suggested_id = NULL WHERE suggested_id = %s", (examiner_id,))
            cur.execute("DELETE FROM examiners WHERE examiner_id = %s", (examiner_id,))
        conn.commit()
        alert_message = "Examiner successfully deleted."
        alert_success = True
    except Exception as e:
        conn.rollback()
        alert_message = f"Error deleting examiner: {str(e)}"
        alert_success = False
        
    if alert_success:
        html = f"""
        <div class="card" style="margin-bottom: 0;">
            <div class="alert alert-success" style="padding: 10px 14px; font-size: 0.85rem; margin-bottom: 0;">
                {alert_message}
            </div>
            <div style="padding: 40px 20px; text-align: center; color: var(--text-muted); border: 1px dashed var(--border); border-radius: var(--radius-md); background-color: rgba(255,255,255,0.01); margin-top: 20px;">
                <div style="font-size: 2rem; margin-bottom: 12px;">👤</div>
                <div style="font-weight: 500; color: var(--text-secondary);">No Person Selected</div>
                <p style="font-size: 0.85rem; margin-top: 4px;">Click any row on the left to view and edit details.</p>
            </div>
        </div>
        """
        return HTMLResponse(content=html, headers={"HX-Trigger": "refreshList"})
    else:
        return HTMLResponse(content=f'<div class="alert alert-danger">{alert_message}</div>')


# ===========================================================================
# TAB 3: MATCHING AND DIAGNOSTICS
# ===========================================================================
def get_top_matching_categories(student, examiner, cats):
    if not student or not examiner:
        return []
    comp_ranks = {'can': 0, 'could': 1, 'cannot': 2}
    matches = []
    for cid, raw_weight in student.categories.items():
        cat = cats.get(cid)
        if not cat:
            continue
        weight = raw_weight if (raw_weight is not None and raw_weight > 0) else cat.default_weight
        lvl = examiner.competences.get(cid, 'cannot')
        matches.append({
            'name': cat.name,
            'competence': lvl,
            'weight': f"{weight:g}"
        })
    matches.sort(key=lambda m: (comp_ranks.get(m['competence'], 2), -float(m['weight']), m['name'].lower()))
    return matches[:3]


def fetch_tab3_data(alert_message: str = "", alert_success: bool = True):
    conn = get_db_conn()
    with conn, conn.cursor(cursor_factory=RealDictCursor) as cur:
        # Fetch examiners dropdown contents
        cur.execute("""
            SELECT ex.examiner_id, ex.full_name, ex.examiner_type, ex.has_dclinpsy, COALESCE(el.limit_n, 3) AS limit_n,
                   (SELECT COUNT(*) FROM examiner_assignments ea WHERE ea.examiner_id = ex.examiner_id) AS assigned_count
            FROM examiners ex
            LEFT JOIN examiner_limits el USING (examiner_id)
            ORDER BY ex.full_name;
        """)
        all_examiners = cur.fetchall()
        internal_examiners = [e for e in all_examiners if e['examiner_type'] == 'internal']
        external_examiners = [e for e in all_examiners if e['examiner_type'] == 'external']
        
        # Load diagnostics data
        cur_raw = conn.cursor()
        cats = diagnostics.fetch_categories(cur_raw)
        students = diagnostics.fetch_students(cur_raw)
        examiners = diagnostics.fetch_examiners(cur_raw)
        assignments = diagnostics.fetch_assignments(cur_raw)
        
        # Check student lists to map remaining metadata
        cur.execute("SELECT student_id, full_name, email, project_title, COALESCE(deferred, FALSE) AS is_deferred, COALESCE(pre_matched, FALSE) AS is_pre_matched FROM students")
        students_db = {row['student_id']: row for row in cur.fetchall()}
        
        student_rows = []
        for sid, ieid, eeid in assignments:
            db_row = students_db.get(sid, {})
            st_obj = students.get(sid)
            iex_obj = examiners.get(ieid) if ieid else None
            eex_obj = examiners.get(eeid) if eeid else None
            
            if not st_obj:
                continue
            
            # Check if both assigned examiners lack DClinPsy (mandatory constraint violation)
            dclinpsy_missing = False
            if iex_obj and eex_obj:
                dclinpsy_missing = not (iex_obj.has_dclinpsy or eex_obj.has_dclinpsy)
                
            # Check mixed qual/quant expertise coverage recommendation
            mixed_mismatch = False
            qual_id = next((cid for cid, cat in cats.items() if cat.name == 'methodology_qualitative'), None)
            quant_id = next((cid for cid, cat in cats.items() if cat.name == 'methodology_quantitative'), None)
            mixed_id = next((cid for cid, cat in cats.items() if cat.name == 'methodology_mixed_qual_quant'), None)

            # Determine joint mixed level
            if mixed_id and mixed_id in st_obj.categories:
                p_level = diagnostics.get_joint_mixed_competence(iex_obj, eex_obj, qual_id, quant_id)
                overrides = {mixed_id: p_level}
                skip_id = mixed_id
                w_mixed = diagnostics._effective_weight(st_obj.categories.get(mixed_id), cats[mixed_id])
                mixed_pen = w_mixed * diagnostics._penalty(p_level, 3, 10)
                mixed_mismatch = (p_level == 'cannot')
            else:
                overrides = None
                skip_id = None
                mixed_pen = 0.0

            i_score, i_pen = (0.0, 0.0)
            i_pen_non_mixed = 0.0
            if iex_obj:
                i_score, i_pen = diagnostics.examiner_metrics(st_obj, iex_obj, cats, 3, 10, override_competences=overrides)
                i_pen_non_mixed = diagnostics.examiner_metrics(st_obj, iex_obj, cats, 3, 10, skip_mixed_id=skip_id)[1]
            
            e_score, e_pen = (0.0, 0.0)
            e_pen_non_mixed = 0.0
            if eex_obj:
                e_score, e_pen = diagnostics.examiner_metrics(st_obj, eex_obj, cats, 3, 10, override_competences=overrides)
                e_pen_non_mixed = diagnostics.examiner_metrics(st_obj, eex_obj, cats, 3, 10, skip_mixed_id=skip_id)[1]
                
            overall_score = round((i_score + e_score) / 2, 1) if (iex_obj and eex_obj) else (i_score or e_score)
            total_penalty = round(i_pen_non_mixed + e_pen_non_mixed + mixed_pen, 2)
            
            internal_top_3 = get_top_matching_categories(st_obj, iex_obj, cats)
            external_top_3 = get_top_matching_categories(st_obj, eex_obj, cats)
            
            student_rows.append({
                "student_id": sid,
                "student_name": db_row.get("full_name", st_obj.name),
                "is_deferred": db_row.get("is_deferred", False),
                "is_pre_matched": db_row.get("is_pre_matched", False),
                "thesis_title": db_row.get("project_title") or "",
                "internal_id": ieid,
                "external_id": eeid,
                "overall_score": overall_score,
                "total_penalty": total_penalty,
                "internal_top_3": internal_top_3,
                "external_top_3": external_top_3,
                "dclinpsy_missing": dclinpsy_missing,
                "mixed_mismatch": mixed_mismatch
            })
            
        # Examiner workloads
        workload_rows = []
        load = {}
        for _, ieid, eeid in assignments:
            load[ieid] = load.get(ieid, 0) + 1
            load[eeid] = load.get(eeid, 0) + 1
            
        for ex_row in all_examiners:
            eid = ex_row['examiner_id']
            assigned = load.get(eid, 0)
            workload_rows.append({
                "examiner_name": ex_row['full_name'],
                "is_internal": ex_row['examiner_type'] == 'internal',
                "limit_n": ex_row['limit_n'],
                "assigned": assigned,
                "remaining": ex_row['limit_n'] - assigned
            })
            
    conn.close()
    return student_rows, workload_rows, internal_examiners, external_examiners, alert_message, alert_success

@app.get("/matching", response_class=HTMLResponse)
def get_matching(request: Request):
    student_rows, workload_rows, internal_examiners, external_examiners, _, _ = fetch_tab3_data()
    return templates.TemplateResponse(request, "matching_tab.html", {
        "active_tab": "matching",
        "student_rows": student_rows,
        "workload_rows": workload_rows,
        "internal_examiners": internal_examiners,
        "external_examiners": external_examiners
    })

@app.post("/matching/clear-assignments")
async def clear_assignments(request: Request):
    conn = get_db_conn()
    alert_msg = ""
    with conn, conn.cursor() as cur:
        cur.execute("""
            DELETE FROM examiner_assignments 
            WHERE student_id NOT IN (
                SELECT student_id FROM students WHERE pre_matched = TRUE
            )
        """)
        alert_msg = "Successfully cleared all solver-generated match allocations. Pre-matched (locked) assignments were kept."
    conn.commit()
    conn.close()
    
    student_rows, workload_rows, internal_examiners, external_examiners, _, _ = fetch_tab3_data(alert_msg, True)
    return templates.TemplateResponse(request, "diagnostics_tables.html", {
        "student_rows": student_rows,
        "workload_rows": workload_rows,
        "internal_examiners": internal_examiners,
        "external_examiners": external_examiners,
        "alert_message": alert_msg,
        "alert_success": True
    })

@app.post("/matching/run-solver")
async def run_solver(
    weight_could: int = Form(3),
    weight_cannot: int = Form(10),
    weight_group_miss: int = Form(5),
    max_seconds: int = Form(300),
    threads: int = Form(8)
):
    try:
        cmd = [
            sys.executable,
            "matching_solver_2.py",
            "--dsn", DB_DSN,
            "--weight-could", str(weight_could),
            "--weight-cannot", str(weight_cannot),
            "--weight-group-miss", str(weight_group_miss),
            "--max-seconds", str(max_seconds),
            "--threads", str(threads)
        ]
        res = subprocess.run(cmd, capture_output=True, text=True, timeout=max_seconds + 10)
        output = res.stdout + "\n" + res.stderr
        
        # Prepend header
        status_msg = "SUCCESS" if res.returncode == 0 else "FAILED"
        header = f"=== SOLVER EXECUTION {status_msg} (Exit Code: {res.returncode}) ===\n"
        return HTMLResponse(content=f"<pre>{header + output}</pre><div hx-get='/matching/tables' hx-trigger='load' hx-target='#diagnostics-container'></div>")
    except Exception as e:
        return HTMLResponse(content=f"<pre>Exception occurred running solver: {str(e)}</pre>")

@app.get("/matching/tables")
def get_matching_tables(request: Request):
    student_rows, workload_rows, internal_examiners, external_examiners, alert_message, alert_success = fetch_tab3_data()
    return templates.TemplateResponse(request, "diagnostics_tables.html", {
        "student_rows": student_rows,
        "workload_rows": workload_rows,
        "internal_examiners": internal_examiners,
        "external_examiners": external_examiners,
        "alert_message": alert_message,
        "alert_success": alert_success
    })

@app.post("/matching/update-assignment")
async def update_assignment(request: Request, student_id: int = Form(...), role: str = Form(...), examiner_id: Optional[int] = Form(None)):
    conn = get_db_conn()
    alert_msg = ""
    alert_ok = True
    with conn, conn.cursor() as cur:
        if not examiner_id:
            cur.execute("DELETE FROM examiner_assignments WHERE student_id = %s AND role = %s", (student_id, role))
            alert_msg = f"Successfully removed {role} examiner assignment."
        else:
            # Check ban first
            cur.execute("""
                SELECT reason FROM examiner_student_bans 
                WHERE student_id = %s AND examiner_id = %s
            """, (student_id, examiner_id))
            ban_row = cur.fetchone()
            if ban_row:
                alert_ok = False
                alert_msg = f"PAIRING CANNOT BE SAVED: This pairing is BANNED! Reason: {ban_row[0] or 'No reason provided'}."
            else:
                # Check workload limit hard rule
                cur.execute("SELECT COALESCE(limit_n, 3) FROM examiner_limits WHERE examiner_id = %s", (examiner_id,))
                limit_row = cur.fetchone()
                limit_n = limit_row[0] if limit_row else 3
                
                # Check current assignments (excluding this student)
                cur.execute("SELECT COUNT(*) FROM examiner_assignments WHERE examiner_id = %s AND student_id != %s", (examiner_id, student_id))
                curr_assigned = cur.fetchone()[0]
                
                if curr_assigned + 1 > limit_n:
                    alert_ok = False
                    cur.execute("SELECT full_name FROM examiners WHERE examiner_id = %s", (examiner_id,))
                    ex_name = cur.fetchone()[0]
                    alert_msg = f"PAIRING CANNOT BE SAVED: {ex_name} has reached their workload limit of {limit_n} vivas."
                else:
                    cur.execute("DELETE FROM examiner_assignments WHERE student_id = %s AND role = %s", (student_id, role))
                    cur.execute("""
                        INSERT INTO examiner_assignments (student_id, role, examiner_id)
                        VALUES (%s, %s, %s)
                    """, (student_id, role, examiner_id))
                    alert_msg = "Assignment updated successfully."
                
    conn.commit()
    conn.close()
    
    student_rows, workload_rows, internal_examiners, external_examiners, _, _ = fetch_tab3_data(alert_msg, alert_ok)
    return templates.TemplateResponse(request, "diagnostics_tables.html", {
        "student_rows": student_rows,
        "workload_rows": workload_rows,
        "internal_examiners": internal_examiners,
        "external_examiners": external_examiners,
        "alert_message": alert_msg,
        "alert_success": alert_ok
    })

@app.post("/matching/toggle-deferred")
async def toggle_deferred(request: Request, student_id: int = Form(...)):
    conn = get_db_conn()
    alert_msg = ""
    with conn, conn.cursor() as cur:
        cur.execute("UPDATE students SET deferred = NOT COALESCE(deferred, FALSE) WHERE student_id = %s RETURNING deferred", (student_id,))
        new_val = cur.fetchone()[0]
        alert_msg = f"Student deferred status set to {new_val}."
    conn.commit()
    conn.close()
    
    student_rows, workload_rows, internal_examiners, external_examiners, _, _ = fetch_tab3_data(alert_msg, True)
    return templates.TemplateResponse(request, "diagnostics_tables.html", {
        "student_rows": student_rows,
        "workload_rows": workload_rows,
        "internal_examiners": internal_examiners,
        "external_examiners": external_examiners,
        "alert_message": alert_msg,
        "alert_success": True
    })

@app.post("/matching/toggle-prematched")
async def toggle_prematched(request: Request, student_id: int = Form(...)):
    conn = get_db_conn()
    alert_msg = ""
    with conn, conn.cursor() as cur:
        cur.execute("UPDATE students SET pre_matched = NOT COALESCE(pre_matched, FALSE) WHERE student_id = %s RETURNING pre_matched", (student_id,))
        new_val = cur.fetchone()[0]
        
        # Get student's name
        cur.execute("SELECT full_name FROM students WHERE student_id = %s", (student_id,))
        st_name = cur.fetchone()[0]
        
        status_str = "PINNED (pre-matched & excluded from solver)" if new_val else "UNPINNED (included in solver matching pool)"
        alert_msg = f"{st_name} is now {status_str}."
    conn.commit()
    conn.close()
    
    student_rows, workload_rows, internal_examiners, external_examiners, _, _ = fetch_tab3_data(alert_msg, True)
    return templates.TemplateResponse(request, "diagnostics_tables.html", {
        "student_rows": student_rows,
        "workload_rows": workload_rows,
        "internal_examiners": internal_examiners,
        "external_examiners": external_examiners,
        "alert_message": alert_msg,
        "alert_success": True
    })


# ===========================================================================
# TAB 4: MAIL MERGE SETUP
# ===========================================================================
def fetch_tab4_paths():
    conn = get_db_conn()
    paths = {}
    with conn, conn.cursor() as cur:
        cur.execute("SELECT key, value FROM system_settings")
        for k, v in cur.fetchall():
            paths[k] = v
    conn.close()
    return paths

@app.get("/mail-merge", response_class=HTMLResponse)
def get_mail_merge(request: Request):
    paths = fetch_tab4_paths()
    conn = get_db_conn()
    with conn, conn.cursor(cursor_factory=RealDictCursor) as cur:
        cur.execute("SELECT student_id, full_name FROM students ORDER BY full_name")
        students = cur.fetchall()
    conn.close()
    return templates.TemplateResponse(request, "mail_merge_tab.html", {
        "active_tab": "mail-merge",
        "paths": paths,
        "students": students
    })

@app.post("/mail-merge/save-paths")
async def save_paths(
    path_trainee: str = Form(...),
    path_internal: str = Form(...),
    path_triad: str = Form(...)
):
    conn = get_db_conn()
    with conn, conn.cursor() as cur:
        cur.execute("UPDATE system_settings SET value = %s WHERE key = 'path_trainee'", (path_trainee.strip(),))
        cur.execute("UPDATE system_settings SET value = %s WHERE key = 'path_internal'", (path_internal.strip(),))
        cur.execute("UPDATE system_settings SET value = %s WHERE key = 'path_triad'", (path_triad.strip(),))
    conn.commit()
    conn.close()
    return HTMLResponse(content='<div class="alert alert-success">Paths updated successfully in database settings.</div>')

@app.post("/mail-merge/preview")
async def preview_mail(student_id: int = Form(...)):
    paths = fetch_tab4_paths()
    
    conn = get_db_conn()
    with conn, conn.cursor(cursor_factory=RealDictCursor) as cur:
        # Fetch selected student details
        cur.execute("""
            SELECT s.student_id, s.full_name AS student_name, s.project_title,
                   sup.supervisors,
                   MAX(CASE WHEN ea.role='internal' THEN ex.full_name END) AS internal_name,
                   MAX(CASE WHEN ea.role='external' THEN ex.full_name END) AS external_name
            FROM students s
            LEFT JOIN student_supervisors sup USING (student_id)
            LEFT JOIN examiner_assignments ea USING (student_id)
            LEFT JOIN examiners ex ON ea.examiner_id = ex.examiner_id
            WHERE s.student_id = %s
            GROUP BY s.student_id, sup.supervisors
        """, (student_id,))
        row = cur.fetchone()
    conn.close()
    
    if not row:
        return HTMLResponse(content='<div style="color: var(--danger);">Student not found.</div>')
        
    student_name = row['student_name']
    project_title = row['project_title'] or "[No Project Title]"
    supervisors = row['supervisors'] or "[No Supervisors Configured]"
    internal_name = row['internal_name'] or "[No Internal Assigned]"
    external_name = row['external_name'] or "[No External Assigned]"
    
    # Let's perform rendering in-memory for preview
    # 1. Trainee Email
    trainee_text = ""
    try:
        p_path = Path(paths.get('path_trainee', 'templates/Trainee_Initial_Email.docx'))
        if p_path.exists():
            doc = DocxTemplate(str(p_path))
            context = {
                'student': student_name,
                'internal_examiner': internal_name,
                'external_examiner': external_name,
                'project_title': project_title
            }
            doc.render(context)
            temp_f = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
            temp_f.close()
            doc.save(temp_f.name)
            
            # Read paragraphs
            docx_doc = Document(temp_f.name)
            paragraphs = [p.text for p in docx_doc.paragraphs if p.text.strip()]
            trainee_text = "".join(f"<p style='margin-bottom:8px;'>{p}</p>" for p in paragraphs)
            Path(temp_f.name).unlink()
        else:
            trainee_text = f"<em>Trainee Template file not found at: {p_path}</em>"
    except Exception as e:
        trainee_text = f"<em>Error rendering trainee preview: {str(e)}</em>"
        
    # 2. Internal Examiner Invite Email
    internal_text = ""
    try:
        p_path = Path(paths.get('path_internal', 'templates/Internal_Examiner_Initial_Email.docx'))
        if p_path.exists():
            doc = DocxTemplate(str(p_path))
            context = {
                'internal_examiner': internal_name,
                'student': student_name,
                'project_title': project_title,
                'supervisors': supervisors
            }
            doc.render(context)
            temp_f = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
            temp_f.close()
            doc.save(temp_f.name)
            
            docx_doc = Document(temp_f.name)
            paragraphs = [p.text for p in docx_doc.paragraphs if p.text.strip()]
            internal_text = "".join(f"<p style='margin-bottom:8px;'>{p}</p>" for p in paragraphs)
            Path(temp_f.name).unlink()
        else:
            internal_text = f"<em>Internal Template file not found at: {p_path}</em>"
    except Exception as e:
        internal_text = f"<em>Error rendering internal invite preview: {str(e)}</em>"
        
    # 3. Triad Email (External)
    external_text = ""
    try:
        p_path = Path(paths.get('path_triad', 'templates/External_Examiner_Initial_Email.docx'))
        if p_path.exists():
            doc = DocxTemplate(str(p_path))
            context = {
                'external_examiner': external_name,
                'student': student_name,
                'internal_examiner': internal_name,
                'project_title': project_title
            }
            doc.render(context)
            temp_f = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
            temp_f.close()
            doc.save(temp_f.name)
            
            docx_doc = Document(temp_f.name)
            paragraphs = [p.text for p in docx_doc.paragraphs if p.text.strip()]
            external_text = "".join(f"<p style='margin-bottom:8px;'>{p}</p>" for p in paragraphs)
            Path(temp_f.name).unlink()
        else:
            external_text = f"<em>External Template file not found at: {p_path}</em>"
    except Exception as e:
        external_text = f"<em>Error rendering external/triad preview: {str(e)}</em>"

    html = f"""
    <div style="margin-bottom: 24px;">
        <h4 style="border-bottom: 1px solid var(--border); padding-bottom: 4px; margin-bottom: 8px; color: var(--accent);">
            Trainee Email Preview
        </h4>
        <div style="background-color: var(--bg-secondary); padding: 12px; border-radius: 4px; border: 1px solid var(--border); max-height: 200px; overflow-y: auto;">
            {trainee_text}
        </div>
    </div>
    
    <div style="margin-bottom: 24px;">
        <h4 style="border-bottom: 1px solid var(--border); padding-bottom: 4px; margin-bottom: 8px; color: var(--success);">
            Internal Examiner Email Preview
        </h4>
        <div style="background-color: var(--bg-secondary); padding: 12px; border-radius: 4px; border: 1px solid var(--border); max-height: 200px; overflow-y: auto;">
            {internal_text}
        </div>
    </div>

    <div>
        <h4 style="border-bottom: 1px solid var(--border); padding-bottom: 4px; margin-bottom: 8px; color: var(--warning);">
            Triad (External Examiner) Email Preview
        </h4>
        <div style="background-color: var(--bg-secondary); padding: 12px; border-radius: 4px; border: 1px solid var(--border); max-height: 200px; overflow-y: auto;">
            {external_text}
        </div>
    </div>
    """
    return HTMLResponse(content=html)


# Subprocesses to run actual Outlook merge scripts
@app.post("/mail-merge/run/trainee")
def run_trainee_merge():
    try:
        cmd = [sys.executable, "trainee_email_initial.py", "--dsn", DB_DSN]
        res = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
        output = res.stdout + "\n" + res.stderr
        status_msg = "SUCCESS" if res.returncode == 0 else "FAILED"
        return HTMLResponse(content=f"<pre>=== TRAINEE EMAIL MERGE {status_msg} ===\n{output}</pre>")
    except Exception as e:
        return HTMLResponse(content=f"<pre>Exception: {str(e)}</pre>")

@app.post("/mail-merge/run/internal")
def run_internal_merge():
    try:
        cmd = [sys.executable, "intexaminer_email1.py", "--dsn", DB_DSN]
        res = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
        output = res.stdout + "\n" + res.stderr
        status_msg = "SUCCESS" if res.returncode == 0 else "FAILED"
        return HTMLResponse(content=f"<pre>=== INTERNAL EMAIL MERGE {status_msg} ===\n{output}</pre>")
    except Exception as e:
        return HTMLResponse(content=f"<pre>Exception: {str(e)}</pre>")

@app.post("/mail-merge/run/triad")
def run_triad_merge():
    try:
        cmd = [sys.executable, "triad_emailer2.py", "--dsn", DB_DSN]
        res = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
        output = res.stdout + "\n" + res.stderr
        status_msg = "SUCCESS" if res.returncode == 0 else "FAILED"
        return HTMLResponse(content=f"<pre>=== TRIAD EMAIL MERGE {status_msg} ===\n{output}</pre>")
    except Exception as e:
        return HTMLResponse(content=f"<pre>Exception: {str(e)}</pre>")
